import streamlit as st
import os
import re
from app import process_meeting
from io import BytesIO
from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet

# ---------------------------------------------------
# PAGE CONFIG
# ---------------------------------------------------
st.set_page_config(
    page_title="Smart Meeting Minutes Generator",
    layout="centered"
)

st.title("📝 Smart Meeting Minutes Generator")
st.write("Upload meeting audio and get structured meeting minutes.")

# ---------------------------------------------------
# SESSION STATE INITIALIZATION
# ---------------------------------------------------
if "processed" not in st.session_state:
    st.session_state.processed = False

if "overall" not in st.session_state:
    st.session_state.overall = ""

if "speaker_summaries" not in st.session_state:
    st.session_state.speaker_summaries = {}

if "actions" not in st.session_state:
    st.session_state.actions = []

if "metadata" not in st.session_state:
    st.session_state.metadata = {}

if "minutes_text" not in st.session_state:
    st.session_state.minutes_text = ""

if "gemini_error" not in st.session_state:
    st.session_state.gemini_error = None

# ---------------------------------------------------
# FILE UPLOAD
# ---------------------------------------------------
uploaded_file = st.file_uploader(
    "Upload audio file",
    type=["wav", "mp3", "m4a"]
)

# Optional: LLM toggle and API key
use_gemini = st.checkbox(
    "Use Google Gemini LLM for minutes (converts dialogue into formal 'what happened' minutes)",
    value=False,
)
gemini_api_key = None
if use_gemini:
    gemini_api_key = st.text_input(
        "Gemini API Key (or set GEMINI_API_KEY in .env / environment)",
        type="password",
        placeholder="Paste your API key here, or leave blank if set in .env",
        help="Get a key at https://aistudio.google.com/apikey",
    )

# ---------------------------------------------------
# PROCESS BUTTON (BEST PRACTICE)
# ---------------------------------------------------
if uploaded_file is not None:

    if st.button("🚀 Process Meeting Audio"):

        os.makedirs("uploads", exist_ok=True)
        upload_path = os.path.join("uploads", "meeting.wav")

        with open(upload_path, "wb") as f:
            f.write(uploaded_file.getbuffer())

        progress = st.progress(0)
        status = st.empty()

        try:
            with st.spinner("Processing meeting audio..."):
                progress.progress(10)
                status.text("Analyzing audio...")

                overall, speaker_summaries, actions, metadata, minutes_text, gemini_error = process_meeting(
                    upload_path,
                    use_llm=use_gemini,
                    gemini_api_key=gemini_api_key if use_gemini else None,
                )

                # SAVE RESULTS TO SESSION STATE
                st.session_state.overall = overall
                st.session_state.speaker_summaries = speaker_summaries
                st.session_state.actions = actions
                st.session_state.metadata = metadata
                st.session_state.minutes_text = minutes_text
                st.session_state.gemini_error = gemini_error
                st.session_state.processed = True

                progress.progress(100)
                status.text("Processing completed")

            st.success("✅ Meeting minutes generated successfully!")
            if st.session_state.gemini_error:
                st.warning("⚠️ Gemini API failed (rule-based minutes used): " + st.session_state.gemini_error)

        except Exception as e:
            st.error(f"❌ Error: {e}")

# ---------------------------------------------------
# DISPLAY RESULTS (ONLY IF PROCESSED)
# ---------------------------------------------------
if st.session_state.processed:

    # -------- METADATA --------
    st.subheader("📌 Metadata")
    for k, v in st.session_state.metadata.items():
        if isinstance(v, list):
            st.write(f"**{k}:** {', '.join(v) if v else 'Not found'}")
        else:
            st.write(f"**{k}:** {v}")

    # -------- EDIT MISSING METADATA --------
    meta = st.session_state.metadata
    date_missing = meta.get("Date", "Not found") == "Not found"
    time_missing = meta.get("Time", "Not found") == "Not found"
    location_missing = meta.get("Location", "Not found") == "Not found"

    if date_missing or time_missing or location_missing:
        st.subheader("✏️ Edit Missing Metadata")
        st.caption("Some metadata was not detected from the audio. You can enter it manually below.")

        new_date = ""
        new_time = ""
        new_location = ""

        if date_missing:
            new_date = st.text_input("📅 Date", placeholder="e.g. March 23, 2026", key="input_date")
        if time_missing:
            new_time = st.text_input("🕐 Time", placeholder="e.g. 3:00 PM", key="input_time")
        if location_missing:
            new_location = st.text_input("📍 Location", placeholder="e.g. Conference Room B", key="input_location")

        if st.button("✅ Apply Changes"):
            updated = False
            minutes = st.session_state.minutes_text

            if date_missing and new_date.strip():
                st.session_state.metadata["Date"] = new_date.strip()
                minutes = re.sub(
                    r"(?i)^Date:\s*Not Mentioned.*$",
                    f"Date: {new_date.strip()}",
                    minutes,
                    flags=re.MULTILINE,
                )
                updated = True

            if time_missing and new_time.strip():
                st.session_state.metadata["Time"] = new_time.strip()
                minutes = re.sub(
                    r"(?i)^Time:\s*Not Mentioned.*$",
                    f"Time: {new_time.strip()}",
                    minutes,
                    flags=re.MULTILINE,
                )
                updated = True

            if location_missing and new_location.strip():
                st.session_state.metadata["Location"] = new_location.strip()
                minutes = re.sub(
                    r"(?i)^Location:\s*Not Mentioned.*$",
                    f"Location: {new_location.strip()}",
                    minutes,
                    flags=re.MULTILINE,
                )
                updated = True

            if updated:
                st.session_state.minutes_text = minutes
                st.success("✅ Metadata updated! The downloads below now include your changes.")
                st.rerun()
            else:
                st.warning("⚠️ Please fill in at least one field before applying.")

    # -------- STRUCTURED MINUTES --------
    st.subheader("🧾 Structured Meeting Minutes")
    st.text(st.session_state.minutes_text)

    # -------- SPEAKER-WISE SUMMARY (visible) --------
    st.subheader("👥 Speaker-wise Summary")
    for sp, txt in st.session_state.speaker_summaries.items():
        st.markdown(f"**{sp}:** {txt}")

    # ======================================================
    # BUILD SPEAKER-WISE SUMMARY TEXT FOR DOWNLOADS
    # ======================================================
    speaker_lines = []
    if st.session_state.speaker_summaries:
        speaker_lines.append("")
        speaker_lines.append("SPEAKER-WISE SUMMARY:")
        for sp, txt in st.session_state.speaker_summaries.items():
            speaker_lines.append(f"- {sp}: {txt}")
    speaker_summary_text = "\n".join(speaker_lines)

    # ======================================================
    # DOWNLOAD SECTION
    # ======================================================
    st.subheader("⬇️ Download Meeting Minutes")

    # -------- TXT --------
    text_content = (st.session_state.minutes_text or "") + speaker_summary_text

    st.download_button(
        "📄 Download as TXT",
        text_content,
        file_name="meeting_minutes.txt"
    )

    # -------- DOCX --------
    doc = Document()
    doc.add_heading("Meeting Minutes", level=1)
    full_text = (st.session_state.minutes_text or "") + speaker_summary_text
    for line in full_text.splitlines():
        doc.add_paragraph(line)

    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)

    st.download_button(
        "📝 Download as DOCX",
        doc_io,
        file_name="meeting_minutes.docx"
    )

    # -------- PDF --------
    pdf_io = BytesIO()
    styles = getSampleStyleSheet()
    pdf = SimpleDocTemplate(pdf_io)

    story = []
    story.append(Paragraph("<b>Meeting Minutes</b>", styles["Title"]))
    full_text_pdf = (st.session_state.minutes_text or "") + speaker_summary_text
    for line in full_text_pdf.splitlines():
        if line.strip().startswith("- "):
            story.append(Paragraph(line, styles["Normal"]))
        elif line.strip().endswith(":"):
            story.append(Paragraph(f"<b>{line}</b>", styles["Heading2"]))
        else:
            story.append(Paragraph(line, styles["Normal"]))

    pdf.build(story)
    pdf_io.seek(0)

    st.download_button(
        "📑 Download as PDF",
        pdf_io,
        file_name="meeting_minutes.pdf"
    )